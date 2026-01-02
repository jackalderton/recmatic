import io
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

def iter_paragraphs_and_tables(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def replace_placeholders_safe(doc: Document, mapping: dict[str, str]):
    keys = sorted(mapping.keys(), key=len, reverse=True)
    for p in iter_paragraphs_and_tables(doc):
        t = p.text or ""
        replaced = False
        for k in keys:
            v = mapping[k]
            if k in t:
                t = t.replace(k, v)
                replaced = True
        if replaced:
            for r in list(p.runs):
                r.clear()
            p.clear()
            p.add_run(t)

def find_placeholder_paragraph(doc: Document, placeholder: str) -> Paragraph | None:
    for p in iter_paragraphs_and_tables(doc):
        if placeholder in (p.text or ""):
            return p
    return None

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def replace_placeholder_with_lines(doc: Document, placeholder: str, lines: list[str]):
    target = find_placeholder_paragraph(doc, placeholder)
    if target is None:
        raise ValueError(f"Placeholder '{placeholder}' not found in template.")
    if not lines:
        target.clear()
        return
    target.clear()
    target.add_run(lines[0])
    anchor = target
    for line in lines[1:]:
        anchor = insert_paragraph_after(anchor, line)

def build_docx(template_bytes: bytes, meta: dict, lines: list[str]) -> bytes:
    bio = io.BytesIO(template_bytes)
    doc = Document(bio)
    # Use .get with defaults so older templates or missing fields don't crash
    replace_placeholders_safe(doc, {
        "[PAGE]": meta.get("page", ""),
        "[DATE]": meta.get("date", ""),
        "[URL]": meta.get("url", ""),
        "[TITLE]": meta.get("title", ""),
        "[TITLE LENGTH]": str(meta.get("title_len", 0)),
        "[DESCRIPTION]": meta.get("description", ""),
        "DESCRIPTION": meta.get("description", ""),
        "[DESCRIPTION LENGTH]": str(meta.get("description_len", 0)),
        "[AGENCY]": meta.get("agency", ""),
        "[CLIENT NAME]": meta.get("client_name", ""),
    })
    replace_placeholder_with_lines(doc, "[PAGE BODY CONTENT]", lines)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
