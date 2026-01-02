import os
import secrets
from pathlib import Path
from urllib.parse import urlencode, urlparse
from base64 import b64encode

import streamlit as st

import io
import re
import json
import html
import datetime
import pytz
import requests
from bs4 import BeautifulSoup, Tag, NavigableString, Comment, Doctype, ProcessingInstruction
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

# --- Semantic embeddings helpers (local MiniLM) ---
from semantic_embeddings import (
    parse_semantic_queries,
    compute_query_similarity_scores,
    format_semantic_scores,
)

APP_DIR = Path(__file__).resolve().parent
ICON_CANDIDATES = [
    APP_DIR / "assets" / "CrFavicon.png",
    APP_DIR / "CrFavicon.png",
]
icon_path = next((p for p in ICON_CANDIDATES if p.exists()), None)

st.set_page_config(
    page_title="Recmatic - Generate Rec Templates Instantly",
    page_icon=str(icon_path) if icon_path else "🧩",
    layout="wide",
)

# --- Auth state ---
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False
if "session_token" not in st.session_state:
    st.session_state["session_token"] = None

url_token = st.query_params.get("token", None)

if url_token and not st.session_state["session_token"]:
    st.session_state["session_token"] = url_token
    st.session_state["authenticated"] = True
elif url_token and url_token == st.session_state["session_token"]:
    st.session_state["authenticated"] = True

if st.session_state.get("authenticated") and st.session_state.get("session_token"):
    st.query_params["token"] = st.session_state["session_token"]
    qs = urlencode({"token": st.session_state["session_token"]})
    st.markdown(
        f"<script>window.history.replaceState(null, '', '?{qs}');</script>",
        unsafe_allow_html=True
    )

ALWAYS_STRIP = {"style", "noscript", "template"}
INLINE_TAGS = {"a","span","strong","em","b","i","u","s","small","sup","sub","mark","abbr","time","code","var","kbd"}
DEFAULT_EXCLUDE = [
    "header", "footer", "nav",
    ".cookie", ".newsletter",
    "[class*='breadcrumb']",
    "[class*='wishlist']",
    "[class*='simplesearch']",
    "[id*='gallery']",
    "[class*='usp']",
    "[class*='feefo']",
    "[class*='associated-blogs']",
    "[class*='popular']",
    ".sr-main.js-searchpage-content.visible",
    "[class~='sr-main'][class~='js-searchpage-content'][class~='visible']",
    "[class*='js-searchpage-content']",
    "[class*='searchpage-content']",
    ".lmd-map-modal-create.js-lmd-map-modal-map",
]
DATE_FMT = "%d/%m/%Y"

NOISE_SUBSTRINGS = (
    "google tag manager",
    "loading results",
    "load more",
    "updating results",
    "something went wrong",
    "filters",
    "apply filters",
    "clear",
    "sort by",
    "to collect end-user usage analytics",
    "place this code immediately before the closing",
)

def safe_filename(name: str, maxlen: int = 120) -> str:
    name = html.unescape(name)
    name = re.sub(r"[\\/*?\"<>|:£#@!^&+=()\[\]{}]", "", name)
    name = re.sub(r"\s+", " ", name)
    name = name.replace(",", "").strip()
    return (name[:maxlen] or "document").rstrip(". ")

def uk_today_str() -> str:
    tz = pytz.timezone("Europe/London")
    return datetime.datetime.now(tz).strftime(DATE_FMT)

def clean_slug_to_name(slug: str) -> str:
    return slug.replace("-", " ").strip().title()

def fallback_page_name_from_url(url: str) -> str:
    path = urlparse(url).path.strip("/")
    parts = [p for p in path.split("/") if p]
    try:
        i = parts.index("destinations")
        if len(parts) > i + 2:
            return clean_slug_to_name(parts[i + 2])
    except ValueError:
        pass
    return clean_slug_to_name(parts[-1] if parts else (urlparse(url).hostname or "Page"))

def fetch_html(url: str) -> tuple[str, bytes]:
    try:
        resp = requests.get(
            url,
            timeout=30,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ContentRecTool/1.0)"}
        )
        resp.raise_for_status()
        return resp.url, resp.content
    except requests.exceptions.HTTPError as e:
        if e.response is not None and e.response.status_code == 403:
            raise RuntimeError(
                "Recmatic couldn't fetch this page (403 Forbidden).\n\n"
                "This site apparently doesn't like our crawler :(\n\n"
                "Unfortunately, that means you'll need to paste in the HTML manually.\n\n"
                "Don't worry! it's still easy!:\n"
                "1. Open the page in your browser\n"
                "2. Right-click → View page source\n"
                "3. Select all (Ctrl/Cmd+A) and Copy\n"
                "4. Paste everything into the box below."
            )
        raise

def normalise_keep_newlines(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"[ \t]*\n[ \t]*", "\n", s)
    return s

def is_noise(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    return any(sub in t for sub in NOISE_SUBSTRINGS)

def annotate_anchor_text(a: Tag, annotate_links: bool) -> str:
    text = a.get_text(" ", strip=True)
    href = a.get("href", "")
    return f"{text} (→ {href})" if (annotate_links and href) else text

def extract_text_preserve_breaks(node: Tag | NavigableString, annotate_links: bool) -> str:
    if isinstance(node, NavigableString):
        return str(node)
    parts = []
    for child in node.children:
        if isinstance(child, NavigableString):
            parts.append(str(child))
        elif isinstance(child, Tag):
            if child.name == "br":
                parts.append("\n")
            elif child.name == "a":
                parts.append(annotate_anchor_text(child, annotate_links))
            else:
                parts.append(extract_text_preserve_breaks(child, annotate_links))
    return "".join(parts)

def extract_schema_jsonld(soup: BeautifulSoup) -> list[str]:
    blocks = []

    def is_ldjson(tag: Tag) -> bool:
        if not isinstance(tag, Tag) or tag.name != "script":
            return False
        t = (tag.get("type") or "").lower()
        return "ld+json" in t

    for sc in soup.find_all(is_ldjson):
        raw = (sc.string or sc.get_text() or "").strip()
        if not raw:
            continue
        try:
            parsed = json.loads(raw)
            pretty = json.dumps(parsed, indent=2, ensure_ascii=False)
            blocks.append(pretty)
        except Exception:
            blocks.append(raw)

    lines = []
    for i, block in enumerate(blocks):
        if i > 0:
            lines.append("")
        lines.extend(block.splitlines())
    return lines

def extract_signposted_lines_from_body(body: Tag, annotate_links: bool, include_img_src: bool = False) -> list[str]:
    lines = []

    def emit_lines(tag_name: str, text: str):
        if tag_name in {"h2", "h3", "h4", "h5", "h6"}:
            if not lines or lines[-1] != "":
                lines.append("")
        text = normalise_keep_newlines(text)
        segments = text.split("\n")
        for seg in segments:
            seg_stripped = seg.strip()
            if seg_stripped:
                if tag_name == "p" and is_noise(seg_stripped):
                    continue
                lines.append(f"<{tag_name}> {seg_stripped}")
            else:
                lines.append("")

    def emit_img(img_tag: Tag):
        if not isinstance(img_tag, Tag) or img_tag.name != "img":
            return
        alt = (img_tag.get("alt") or "").strip().replace('"', '\\"')
        if include_img_src:
            src = (img_tag.get("src") or "").strip().replace('"', '\\"')
            if src:
                lines.append(f'<img alt="{alt}" src="{src}">')
                return
        lines.append(f'<img alt="{alt}">')

    def handle(tag: Tag):
        name = tag.name
        if name in ALWAYS_STRIP or name == "script":
            return
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            txt = extract_text_preserve_breaks(tag, annotate_links)
            if txt.strip():
                emit_lines(name, txt)
            return
        if name == "p":
            txt = tag.get_text(" ", strip=True)
            if txt.strip():
                emit_lines("p", txt)
            for img in tag.find_all("img"):
                emit_img(img)
            return
        if name in {"ul", "ol"}:
            for li in tag.find_all("li", recursive=False):
                txt = extract_text_preserve_breaks(li, annotate_links)
                if txt.strip():
                    emit_lines("p", txt)
                for img in li.find_all("img"):
                    emit_img(img)
                for sub in li.find_all(["ul", "ol"], recursive=False):
                    for sub_li in sub.find_all("li", recursive=False):
                        sub_txt = extract_text_preserve_breaks(sub_li, annotate_links)
                        if sub_txt.strip():
                            emit_lines("p", sub_txt)
                        for img in sub_li.find_all("img"):
                            emit_img(img)
            return

        buf = []
        def flush_buf():
            if not buf:
                return
            joined = normalise_keep_newlines("".join(buf))
            if joined.strip() and not is_noise(joined):
                emit_lines("p", joined)
            buf.clear()

        for child in tag.children:
            if isinstance(child, (Comment, Doctype, ProcessingInstruction)):
                continue
            if isinstance(child, NavigableString):
                buf.append(str(child))
            elif isinstance(child, Tag):
                if child.name == "br":
                    buf.append("\n")
                elif child.name == "img":
                    flush_buf()
                    emit_img(child)
                elif child.name in INLINE_TAGS:
                    buf.append(extract_text_preserve_breaks(child, annotate_links))
                else:
                    flush_buf()
                    handle(child)
        flush_buf()

    for child in body.children:
        if isinstance(child, (Comment, Doctype, ProcessingInstruction)):
            continue
        if isinstance(child, NavigableString):
            raw = normalise_keep_newlines(str(child))
            if raw.strip() and not is_noise(raw):
                emit_lines("p", raw)
            elif raw == "\n":
                lines.append("")
        elif isinstance(child, Tag):
            if child.name == "img":
                emit_img(child)
            else:
                handle(child)

    deduped = []
    for ln in lines:
        if ln == "":
            if not deduped or deduped[-1] != "":
                deduped.append("")
            continue
        if not deduped or ln != deduped[-1]:
            deduped.append(ln)
    return deduped

def remove_before_first_h1_all_levels(body: Tag) -> None:
    if body is None:
        return
    first_h1 = body.find("h1")
    if first_h1 is None:
        return
    chain = []
    node = first_h1
    while node is not None and node != body:
        chain.append(node)
        node = node.parent
    chain.reverse()
    for child in chain:
        for prev in list(child.previous_siblings):
            try:
                if isinstance(prev, Tag):
                    prev.decompose()
                elif isinstance(prev, NavigableString):
                    prev.extract()
            except Exception:
                continue

def iter_paragraphs_and_tables(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for sec in doc.sections:
        headers = [
            getattr(sec, "header", None),
            getattr(sec, "first_page_header", None),
            getattr(sec, "even_page_header", None),
        ]
        footers = [
            getattr(sec, "footer", None),
            getattr(sec, "first_page_footer", None),
            getattr(sec, "even_page_footer", None),
        ]
        for part in [h for h in headers if h] + [f for f in footers if f]:
            for p in part.paragraphs:
                yield p
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

def replace_placeholders_safe(doc: Document, mapping: dict[str, str]):
    keys = sorted(mapping.keys(), key=len, reverse=True)
    for p in iter_paragraphs_and_tables(doc):
        t = p.text or ""
        replaced = False
        for k in keys:
            v = mapping[k]
            if k in t:
                t = t.replace(k, v)
                replaced = True
        if replaced:
            for r in list(p.runs):
                r.clear()
            p.clear()
            p.add_run(t)

def find_placeholder_paragraph(doc: Document, placeholder: str) -> Paragraph | None:
    for p in iter_paragraphs_and_tables(doc):
        if placeholder in (p.text or ""):
            return p
    return None

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def replace_placeholder_with_lines(doc: Document, placeholder: str, lines: list[str]) -> bool:
    """
    Replace a placeholder with multiple lines.
    Returns True if the placeholder existed, False if not.

    Missing placeholders are treated as optional.
    """
    target = find_placeholder_paragraph(doc, placeholder)
    if target is None:
        return False  # <-- missing but not an error

    if not lines:
        target.clear()
        return True

    target.clear()
    target.add_run(lines[0])
    anchor = target
    for line in lines[1:]:
        anchor = insert_paragraph_after(anchor, line)

    return True

def build_docx(template_bytes: bytes, meta: dict, lines: list[str]) -> tuple[bytes, list[str]]:
    bio = io.BytesIO(template_bytes)
    doc = Document(bio)

    # Simple placeholder text replacements
    replace_placeholders_safe(doc, {
        "[PAGE]": meta.get("page", ""),
        "[DATE]": meta.get("date", ""),
        "[URL]": meta.get("url", ""),
        "[TITLE]": meta.get("title", ""),
        "[TITLE LENGTH]": str(meta.get("title_len", 0)),
        "[DESCRIPTION]": meta.get("description", ""),
        "DESCRIPTION": meta.get("description", ""),
        "[DESCRIPTION LENGTH]": str(meta.get("description_len", 0)),
        "[AGENCY]": meta.get("agency", ""),
        "[CLIENT NAME]": meta.get("client_name", ""),
        "[KEYWORDS]": meta.get("keywords", ""),
        # Always remove placeholder unless populated
        "[SEMANTIC SCORES]": meta.get("semantic_scores", ""),
    })

    missing = []

    # PAGE BODY CONTENT
    ok = replace_placeholder_with_lines(doc, "[PAGE BODY CONTENT]", lines)
    if not ok:
        missing.append("[PAGE BODY CONTENT]")

    # SCHEMA
    if meta.get("schema_lines"):
        ok = replace_placeholder_with_lines(doc, "[SCHEMA]", meta["schema_lines"])
        if not ok:
            missing.append("[SCHEMA]")
    else:
        ok = replace_placeholder_with_lines(doc, "[SCHEMA]", [])
        if not ok:
            pass

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    return out.read(), missing

def first_h1_text(soup: BeautifulSoup) -> str | None:
    if not soup.body:
        return None
    h1 = soup.body.find("h1")
    if not h1:
        return None
    txt = extract_text_preserve_breaks(h1, annotate_links=False)
    txt = normalise_keep_newlines(txt)
    txt = re.sub(r"\s+", " ", txt)
    return txt.strip() or None

def process_url(
    url: str,
    exclude_selectors: list[str],
    annotate_links: bool = False,
    remove_before_h1: bool = False,
    include_img_src: bool = False,
):
    final_url, html_bytes = fetch_html(url)
    soup = BeautifulSoup(html_bytes, "lxml")
    schema_lines = extract_schema_jsonld(soup)
    for el in soup.find_all(list(ALWAYS_STRIP)):
        el.decompose()
    body = soup.body or soup
    for sel in exclude_selectors:
        try:
            for el in body.select(sel):
                el.decompose()
        except Exception:
            pass
    try:
        for el in body.find_all(lambda t: isinstance(t, Tag) and t.has_attr('class') and {'sr-main','js-searchpage-content','visible'}.issubset(set(t.get('class', [])))):
            el.decompose()
    except Exception:
        pass
    for sel in [
        '.sr-main.js-searchpage-content.visible',
        "[class~='sr-main'][class~='js-searchpage-content'][class~='visible']",
        "[class*='js-searchpage-content']",
        "[class*='searchpage-content']",
        ".lmd-map-modal-create.js-lmd-map-modal-map",
    ]:
        try:
            for el in body.select(sel):
                el.decompose()
        except Exception:
            pass
    if remove_before_h1:
        remove_before_first_h1_all_levels(body)
    lines = extract_signposted_lines_from_body(body, annotate_links=annotate_links, include_img_src=include_img_src)
    head = soup.head or soup
    title = head.title.string.strip() if (head and head.title and head.title.string) else "N/A"
    meta_el = head.find("meta", attrs={"name": "description"}) if head else None
    description = meta_el.get("content").strip() if (meta_el and meta_el.get("content")) else "N/A"
    page_name = first_h1_text(soup) or fallback_page_name_from_url(final_url)
    meta = {
        "page": page_name,
        "date": uk_today_str(),
        "url": final_url,
        "title": title,
        "title_len": len(title) if title != "N/A" else 0,
        "description": description,
        "description_len": len(description) if description != "N/A" else 0,
        "schema_lines": schema_lines,
    }
    return meta, lines

def process_html_string(
    html_text: str,
    url: str,
    exclude_selectors: list[str],
    annotate_links: bool,
    remove_before_h1: bool,
    include_img_src: bool,
):
    html_bytes = html_text.encode("utf-8", errors="ignore")
    soup = BeautifulSoup(html_bytes, "lxml")
    schema_lines = extract_schema_jsonld(soup)
    for el in soup.find_all(list(ALWAYS_STRIP)):
        el.decompose()
    body = soup.body or soup
    for sel in exclude_selectors:
        try:
            for el in body.select(sel):
                el.decompose()
        except Exception:
            pass
    for sel in [
        '.sr-main.js-searchpage-content.visible',
        "[class~='sr-main'][class~='js-searchpage-content'][class~='visible']",
        "[class*='js-searchpage-content']",
        "[class*='searchpage-content']",
        ".lmd-map-modal-create.js-lmd-map-modal-map",
    ]:
        try:
            for el in body.select(sel):
                el.decompose()
        except Exception:
            pass
    if remove_before_h1:
        remove_before_first_h1_all_levels(body)
    lines = extract_signposted_lines_from_body(body, annotate_links=annotate_links, include_img_src=include_img_src)
    head = soup.head or soup
    title = head.title.string.strip() if (head and head.title and head.title.string) else "N/A"
    meta_el = head.find("meta", attrs={"name": "description"}) if head else None
    description = meta_el.get("content").strip() if (meta_el and meta_el.get("content")) else "N/A"
    page_name = first_h1_text(soup) or fallback_page_name_from_url(url)
    meta = {
        "page": page_name,
        "date": uk_today_str(),
        "url": url,
        "title": title,
        "title_len": len(title) if title != "N/A" else 0,
        "description": description,
        "description_len": len(description) if description != "N/A" else 0,
        "schema_lines": schema_lines,
    }
    return meta, lines

# --- Fonts (local Circular) ---
CANDIDATES = {
    400: [
        APP_DIR / "assets" / "fonts" / "lineto-circular-book.woff2",
        APP_DIR / "assets" / "fonts" / "CircularStd-Book.woff2",
        APP_DIR / "assets" / "lineto-circular-book.woff2",
        APP_DIR / "lineto-circular-book.woff2",
    ],
    700: [
        APP_DIR / "assets" / "fonts" / "lineto-circular-bold.woff2",
        APP_DIR / "assets" / "fonts" / "lineto-circular-bold.woff2",
        APP_DIR / "lineto-circular-bold.woff2",
    ],
}

faces_css = []
have_weight = set()
for weight, paths in CANDIDATES.items():
    p = next((x for x in paths if x.exists()), None)
    if p:
        data = b64encode(p.read_bytes()).decode("utf-8")
        faces_css.append(
            "@font-face {\n"
            "  font-family: 'CircularLocal';\n"
            f"  src: url(data:font/woff2;charset=utf-8;base64,{data}) format('woff2');\n"
            f"  font-weight: {weight};\n"
            "  font-style: normal;\n"
            "  font-display: swap;\n"
            "}\n"
        )
        have_weight.add(weight)

if faces_css:
    st.markdown("<style>\n" + "".join(faces_css) + "</style>", unsafe_allow_html=True)
else:
    st.info("Add Circular WOFF2 fonts to enable local typography.")

force_bold_line = ""
if 400 not in have_weight and 700 in have_weight:
    force_bold_line = "  font-weight: 700;\n"

st.markdown(
    "<style>\n"
    "html, body, [data-testid=\"stAppViewContainer\"] *"
    ":not(.material-icons):not(.material-icons-outlined)"
    ":not(.material-symbols-outlined):not(.material-symbols-rounded):not(.material-symbols-sharp) {\n"
    "  font-family: 'CircularLocal', system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif;\n"
    + force_bold_line +
    "}\n"
    "[data-testid=\"stIconMaterial\"] { display: none !important; }\n"
    "[data-testid=\"stAppViewContainer\"] { background-color: #000000; }\n"
    "[data-testid=\"stAppViewContainer\"] .block-container { background-color: transparent; }\n"
    "section[tabindex=\"0\"] h1:first-of-type { color: #537DFC !important; }\n"
    "h1, h2, h3 { color: #537DFC !important; }\n"
    "[data-testid=\"stSidebar\"] { background-color: #1a1e24; border-right: 1px solid #4A90E2; min-width: 450px; }\n"
    "[data-testid=\"stExpander\"] [data-testid=\"stExpanderHeader\"] {\n"
    "  background-color: #363945; border-radius: 8px; padding: 10px 15px; margin-bottom: 10px;\n"
    "  font-weight: bold; color: #E0E0E0; }\n"
    ".stButton > button { width: 100%; background-color: #323640; color: #E0E0E0;\n"
    "  border: 1px solid #4A90E2; border-radius: 8px; padding: 10px; }\n"
    ".stButton > button:hover { background-color: #4A90E2; color: #fff; border-color: #fff; }\n"
    "[data-testid=\"stTabs\"] button[role=\"tab\"] { background-color: #323640; color: #E0E0E0; }\n"
    "[data-testid=\"stTabs\"] button[role=\"tab\"][aria-selected=\"true\"] { color: #4A90E2; box-shadow: inset 0 -3px 0 0 #4A90E2; }\n"
    "</style>\n",
    unsafe_allow_html=True,
)

st.title("Recmatic")

if "single_docx" not in st.session_state:
    st.session_state.single_docx = None
    st.session_state.single_docx_name = None

# Fallback state
if "run_fallback_preview" not in st.session_state:
    st.session_state["run_fallback_preview"] = False
if "run_fallback_doc" not in st.session_state:
    st.session_state["run_fallback_doc"] = False
if "show_fallback" not in st.session_state:
    st.session_state["show_fallback"] = False
if "pasted_html" not in st.session_state:
    st.session_state["pasted_html"] = ""

# --- Sidebar ---
with st.sidebar:
    if st.session_state.get("authenticated"):
        if st.button("Logout"):
            st.session_state["authenticated"] = False
            st.session_state["session_token"] = None
            st.query_params.clear()
            st.rerun()

    st.link_button("Request a feature", "https://docs.google.com/forms/d/e/1FAIpQLSc8xivp45me2u6dRjBNmBzd7uX_B-qBYOcPXjWZo1sNcbJFzA/viewform?usp=header")

    st.header("Template & Options")

    tpl_file = st.file_uploader(
        "Optional: upload a custom template (.docx)",
        type=["docx"],
    )
    st.caption(
        "You can just leave this empty now! If you have your own template, you can upload that here - Recmatic will use yours instead."
    )

    st.divider()
    st.header("Need a template?")

    TEMPLATE_CANDIDATES = [
        APP_DIR / "assets" / "blank_template.docx",
        APP_DIR / "blank_template.docx",
    ]
    template_path = next((p for p in TEMPLATE_CANDIDATES if p.exists()), None)

    if template_path:
        with open(template_path, "rb") as file:
            st.download_button(
                label="Download the default blank template",
                data=file,
                file_name="blank_template.docx",
                mime=(
                    "application/"
                    "vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )
        st.caption(
            "Hate the template I spent ages making? Cool! You can download and customise this blank one."
        )
    else:
        st.info(
            "Default template file not found. "
            "You can still upload your own .docx template above."
        )

    st.divider()
    st.header("Exclude Selectors")

    exclude_txt = st.text_area(
        "Comma-separated CSS selectors to remove",
        value=", ".join(DEFAULT_EXCLUDE),
        height=120,
    )
    exclude_selectors = [s.strip() for s in exclude_txt.split(",") if s.strip()]

    st.divider()
    st.header("Extra Settings")

    include_schema = st.toggle("Include Schema", value=False)
    annotate_links = st.toggle("Append (→ URL) after anchor text", value=False)
    remove_before_h1 = st.toggle("Delete everything before first <h1>", value=False)
    include_img_src = st.toggle("Include image sources", value=False)

    # --- Semantic Embeddings (BETA) ---
    st.divider()
    st.header("Semantic Embeddings (BETA)")

    enable_semantic = st.toggle("Toggle semantic embeddings (BETA)", value=False)

    semantic_queries_raw = ""
    if enable_semantic:
        semantic_queries_raw = st.text_input(
            "Semantic queries (comma-separated)",
            value="",
            placeholder="e.g. what would a user search to find this page?",
        )
        st.caption("Scores each query against the full extracted page content (one vector).")

    st.divider()

    try:
        with open("VERSION.txt", "r") as f:
            version = f.read().strip()
        st.caption(f"Version: `{version}`")
    except Exception:
        st.caption("Version: dev")

# --- Keywords UI ---
with st.expander("Add Keywords (Optional)", expanded=False):
    st.caption("This works as intended, but the UI may feel clunky due to re-rendering. No further development is planned.")

    default_row_count = len(st.session_state.get("keywords_list", [])) or 1
    row_count = st.number_input(
        "How many keywords? (Max 10)",
        min_value=1,
        max_value=10,
        step=1,
        value=default_row_count
    )

    existing = st.session_state.get("keywords_list", [])
    while len(existing) < row_count:
        existing.append({"keyword": "", "volume": ""})
    if len(existing) > row_count:
        existing = existing[:row_count]
    st.session_state.keywords_list = existing

    new_keywords = []
    for idx, pair in enumerate(st.session_state.keywords_list):
        col1, col2 = st.columns([3, 1])
        with col1:
            kw = st.text_input("", value=pair["keyword"], key=f"kw_{idx}", placeholder="Keyword")
        with col2:
            vol = st.text_input("", value=pair["volume"], key=f"vol_{idx}", placeholder="Vol")
        new_keywords.append({"keyword": kw, "volume": vol})

    st.session_state.keywords_list = new_keywords

def parse_volume(vol) -> int:
    s = str(vol).strip().lower().replace(",", "")
    mult = 1
    if s.endswith("k"):
        mult = 1000
        s = s[:-1]
    elif s.endswith("m"):
        mult = 1_000_000
        s = s[:-1]
    try:
        return int(float(s) * mult)
    except Exception:
        return 0

sorted_keywords = sorted(
    [
        item for item in st.session_state.keywords_list
        if item['keyword'].strip() and item['volume'].strip()
    ],
    key=lambda x: parse_volume(x['volume']),
    reverse=True,
)

formatted_keywords = ", ".join(
    f"{item['keyword']} ({item['volume']})"
    for item in sorted_keywords
)

col0a, col0b = st.columns([1, 1])
with col0a:
    client_name = st.text_input("Client Name", value="", placeholder="e.g., LeShuttle")
with col0b:
    agency_name = st.text_input("Agency/Practitioner Name", value="", placeholder="e.g., Crafted")

url = st.text_input("URL", value="", placeholder="https://www.example.com")

do_doc = st.button("Generate!")

# ── Main crawl path ──────────────────────────────────────────────────────────────
if do_doc:
    if not url.strip():
        st.error("Please enter a URL.")
    else:
        try:
            meta, lines = process_url(
                url,
                exclude_selectors,
                annotate_links=annotate_links,
                remove_before_h1=remove_before_h1,
                include_img_src=include_img_src,
            )

            meta["agency"] = agency_name.strip()
            meta["client_name"] = client_name.strip()
            meta["keywords"] = formatted_keywords

            # Always remove placeholder unless explicitly populated
            meta["semantic_scores"] = ""

            if enable_semantic and semantic_queries_raw.strip():
                queries = parse_semantic_queries(semantic_queries_raw)
                if queries:
                    content_text = "\n".join(lines)
                    scored = compute_query_similarity_scores(content_text, queries)
                    meta["semantic_scores"] = format_semantic_scores(scored)

            if not include_schema:
                meta["schema_lines"] = []

            st.success("Extracted successfully.")

            with st.expander("Meta (preview)", expanded=True):
                st.write(meta)

            with st.expander("Signposted content (preview)", expanded=False):
                st.text("\n".join(lines))

            with st.expander("Schema (preview)", expanded=False):
                schema_preview = "\n".join(meta.get("schema_lines", [])) or "No JSON-LD schema found."
                st.text(schema_preview)

            tpl_bytes = None
            if tpl_file is not None:
                tpl_bytes = tpl_file.read()
            elif template_path is not None:
                with open(template_path, "rb") as f:
                    tpl_bytes = f.read()

            if tpl_bytes is None:
                st.error(
                    "We couldn't find any template file. Either upload one in the "
                    "sidebar or make sure blank_template.docx exists in /assets "
                    "or the app folder."
                )
            else:
                out_bytes, missing = build_docx(tpl_bytes, meta, lines)
                raw_title = meta.get("page", "Untitled Page")
                cleaned_title = safe_filename(raw_title)
                client_suffix = ""
                if meta.get("client_name"):
                    client_clean = safe_filename(meta["client_name"])
                    client_suffix = f"_{client_clean}"
                fname = f"{cleaned_title} - Content Recommendations{client_suffix}.docx"

                st.session_state.single_docx = out_bytes
                st.session_state.single_docx_name = fname

                if missing:
                    missing_str = ", ".join(missing)
                    st.warning(
                        "Just FYI – your template does not contain: "
                        f"{missing_str}. Those sections were skipped. "
                        "If that's intentional, you're all set."
                    )

            st.session_state["show_fallback"] = False

        except RuntimeError as e:
            st.error(str(e))
            st.session_state["show_fallback"] = True

        except Exception as e:
            st.exception(e)

# ── Fallback: Paste HTML Instead ────────────────────────────────────────────────
if st.session_state.get("show_fallback"):
    with st.expander("Paste HTML Instead", expanded=True):
        st.session_state["pasted_html"] = st.text_area(
            "Paste full HTML source here",
            value=st.session_state.get("pasted_html", ""),
            height=350,
        )

        pasted_html = st.session_state["pasted_html"]

        colx, coly = st.columns([1, 1])
        with colx:
            if st.button("Process pasted HTML (preview only)"):
                st.session_state["run_fallback_preview"] = True
        with coly:
            if st.button("Generate DOCX from pasted HTML"):
                st.session_state["run_fallback_doc"] = True

    pasted_html = st.session_state.get("pasted_html", "")

    if st.session_state.get("run_fallback_preview") and pasted_html.strip():
        st.session_state["run_fallback_preview"] = False

        meta, lines = process_html_string(
            pasted_html,
            url,
            exclude_selectors,
            annotate_links,
            remove_before_h1,
            include_img_src,
        )

        meta["agency"] = agency_name.strip()
        meta["client_name"] = client_name.strip()
        meta["keywords"] = formatted_keywords

        # Keep semantic placeholder empty for preview-only path
        meta["semantic_scores"] = ""

        if not include_schema:
            meta["schema_lines"] = []

        st.success("HTML processed successfully.")

        with st.expander("Meta (preview)", expanded=True):
            st.write(meta)

        with st.expander("Signposted content (preview)", expanded=False):
            st.text("\n".join(lines))

        with st.expander("Schema (preview)", expanded=False):
            schema_preview = "\n".join(meta.get("schema_lines", [])) or "No JSON-LD schema found."
            st.text(schema_preview)

    if st.session_state.get("run_fallback_doc") and pasted_html.strip():
        st.session_state["run_fallback_doc"] = False

        meta, lines = process_html_string(
            pasted_html,
            url,
            exclude_selectors,
            annotate_links,
            remove_before_h1,
            include_img_src,
        )

        meta["agency"] = agency_name.strip()
        meta["client_name"] = client_name.strip()
        meta["keywords"] = formatted_keywords

        # Always remove placeholder unless explicitly populated
        meta["semantic_scores"] = ""

        if enable_semantic and semantic_queries_raw.strip():
            queries = parse_semantic_queries(semantic_queries_raw)
            if queries:
                content_text = "\n".join(lines)
                scored = compute_query_similarity_scores(content_text, queries)
                meta["semantic_scores"] = format_semantic_scores(scored)

        if not include_schema:
            meta["schema_lines"] = []

        tpl_bytes = None
        if tpl_file is not None:
            tpl_bytes = tpl_file.read()
        elif template_path is not None:
            with open(template_path, "rb") as f:
                tpl_bytes = f.read()

        if tpl_bytes is None:
            st.error(
                "We couldn't find any template file. Either upload one in the "
                "sidebar or make sure blank_template.docx exists in /assets "
                "or the app folder."
            )
        else:
            out_bytes, missing = build_docx(tpl_bytes, meta, lines)
            raw_title = meta.get("page", "Untitled Page")
            cleaned_title = safe_filename(raw_title)

            client_suffix = ""
            if meta.get("client_name"):
                client_clean = safe_filename(meta["client_name"])
                client_suffix = f"_{client_clean}"

            fname = f"{cleaned_title} - Content Recommendations{client_suffix}.docx"

            st.session_state.single_docx = out_bytes
            st.session_state.single_docx_name = fname

            if missing:
                missing_str = ", ".join(missing)
                st.warning(
                    "Just FYI – your template does not contain: "
                    f"{missing_str}. Those sections were skipped. "
                    "If that's intentional, you're all set."
                )

if st.session_state.single_docx:
    st.download_button(
        "Download DOCX",
        data=st.session_state.single_docx,
        file_name=st.session_state.single_docx_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="dl_single_docx",
    )
